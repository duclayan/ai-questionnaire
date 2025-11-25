from django.shortcuts import render
from rest_framework import viewsets
from rest_framework.decorators import api_view
from rest_framework.response import Response

from .models import *
from .serializer import *
from rest_framework.views import APIView
from rest_framework import status
from docx.shared import Inches
from openai import AzureOpenAI
from dotenv import load_dotenv
from io import BytesIO
from docx import Document
import os
import httpx
import requests
from django.http import HttpResponse

from django.contrib.auth import authenticate
from rest_framework.authtoken.models import Token
from rest_framework.response import Response
from rest_framework.permissions import AllowAny

from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.authentication import JWTAuthentication

from rest_framework import status
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
import base64
import logging
import json 
from rest_framework.parsers import MultiPartParser
from django.shortcuts import get_object_or_404
import os
from django.conf import settings
logger = logging.getLogger(__name__)
class MockRequest:
    def __init__(self, data):
        self.data = data
from rest_framework.parsers import MultiPartParser, FormParser
import os
import docx
import PyPDF2
import tempfile
from django.conf import settings
from django.db import IntegrityError
from pathlib import Path
import asyncio
import uuid
from .transcribe import transcribe_audio
from .diagram_service import generate_architecture_diagram
class ClassicGPT:
    def getGPTResponse(self, prompt):

        #prepare the prompt
        AZURE_OPENAI_ENDPOINT= os.getenv("AZURE_OPENAI_ENDPOINT")
        AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
        AZURE_OPENAI_VERSION = os.getenv("LATEST_AZURE_OPENAI_VERSION")

        #CALL GPT
        # Initialize the Azure OpenAI client

        if os.getenv("DJANGO_DEVELOPMENT", "False") == "True":
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}",
                http_client=httpx.Client(verify=False)
            )
        else:
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}"
            )

        # Create a chat completion
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                    {"role": "system", "content": "You are a helpful assistant that responds in Markdown. Help me with my math homework!"},
                    {"role": "user", "content": [
                        {"type": "text", "text":f"{prompt}"},
                    ]}
                ],
            temperature=0.0,
        )
        # Extract the generated text from the response
        generated_text = response.choices[0].message.content
        print("Generated Text Answer:", generated_text)
        # Return a JSON response
        return generated_text

class DocumentGenerator:
    def create_document_response(self, gpt_response):
        doc = Document()

        doc.add_heading('Report Summary', level=0)  # Main heading
        
        # Split the GPT response into sections based on new lines
        sections = gpt_response.strip().split("\n\n")  # Split by double newlines for sections

        for section in sections:
            if section:  
                # Ensure the section is not empty
                # Split the section into lines
                lines = section.splitlines()
                if lines:
                    # The first line is treated as the heading
                    heading = lines[0].strip()
                    if heading.startswith("### "):
                        heading = heading[4:].strip()  # Remove '### ' prefix
                        if (heading == 'Counter Measure'):
                            doc.add_heading(heading, level=3)  
                        else:
                            doc.add_heading(heading, level=1)  
                    else:
                        body = "\n".join(line.strip() for line in lines if line.strip())
                        if body:
                            doc.add_paragraph(body)  # Add the body text below the heading

        # Add Architecture Diagram
        # Add the saved diagram image to the document
        self.add_diagram_to_document(doc)

        # Save the document to a BytesIO object
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Create the HTTP response with the document
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=report_summary.docx'
        return response
    def add_diagram_to_document(self, doc):
        """Add saved diagram image to Word document."""
        file_name = 'diagram.png'
        try:
            # Assuming your diagram is saved as 'diagram.png' in MEDIA_ROOT
            file_path = os.path.join(default_storage.location, file_name)
            if os.path.exists(file_path):
                # Add picture to document with specified width
                # Adjust the width as needed (e.g., 6.0 inches)
                doc.add_heading("Architecture Diagram", level=1)  
                doc.add_picture(file_path, width=Inches(6.0))  # Set width to 6 inches
                default_storage.delete(file_name)
            else:
                print("Diagram image not found.")
        except Exception as e:
            print(f"Error adding diagram to document: {e}")
class UserView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = LoginSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=400)

        username = serializer.validated_data['username']
        password = serializer.validated_data['password']
        
        logger.info(f"Received login attempt for username: {username}")

        if not username or not password:
            return Response({'error': 'Please provide both username and password'}, status=400)

        # Ensure password is a string
        if isinstance(password, (list, tuple)):
            password = password[0] if password else ''
        password = str(password)

        logger.debug(f"Password type after conversion: {type(password)}")

        user = authenticate(request, username=username, password=password)
        
        if user:
            token, created = Token.objects.get_or_create(user=user)
            return Response({'token': token.key})
        else:
            logger.warning(f"Failed login attempt for username: {username}")
            return Response({'error': 'Invalid credentials'}, status=401)
class QuestionListView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    def get(self, request):
        current_category = request.query_params.get("currentCategory")
        if current_category == 'all':
            questions = Question.objects.all()
        elif current_category and current_category != "Report":
            questions = Question.objects.filter(category=current_category)
        else:
            questions = Question.objects.none()

        serializer = QuestionSerializer(questions, many=True)
        return Response({"question_list": serializer.data}, status=status.HTTP_200_OK)
    def post(self, request):
        serializer = QuestionSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
class openAICleanVersion(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    def post(self, request):
        # Get user input
        data = request.data.get("text")
        # Load environment variables
        load_dotenv()
        AZURE_OPENAI_ENDPOINT= os.getenv("AZURE_OPENAI_ENDPOINT")
        AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
        AZURE_OPENAI_VERSION = os.getenv("LATEST_AZURE_OPENAI_VERSION")

        # Initialize the Azure OpenAI client
        if os.getenv("DJANGO_DEVELOPMENT", "False") == "True":
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}",
                http_client=httpx.Client(verify=False)
            )
        else:
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}"
            )

        # Create a chat completion
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            # removing temperature and max_completion changed to max_completion_tokens
            # this is to accommodate to the new change of azureopenai error saying this method
            # is depreciated
            # Removed the max_completion completely when using O1 Models
            # temperature=0.9,
            # max_completion_tokens=2000,
          messages=[{ 
                        "role": "user",
                        "content": f"""
                          {data}
                        """,
                    }]
        )
        # Extract the generated text from the response
        generated_text = response.choices[0].message.content

        # Return a JSON response
        return Response({"generated_text": generated_text})
class openAICleanVersion_O4MINI(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    def _build_prompt(self, data):
        """Build the appropriate prompt based on available data"""
        return f"""
        Strictly return only mermaidjs code in plaintext. No explanations or other texts included
        Language: english
        Also consider architecture-beta of mermaidjs. Sample:
        architecture-beta
            group api(cloud)[API]

            service db(database)[Database] in api
            service disk1(disk)[Storage] in api
            service disk2(disk)[Storage] in api
            service server(server)[Server] in api

            db:L -- R:server
            disk1:T -- B:server
            disk2:T -- B:db
        ---
            {data}
        """


    def post(self, request):
        # Get user input
        data = request.data.get("text")
        final_prompt =self._build_prompt(data)
        # Load environment variables
        load_dotenv()
        AZURE_OPENAI_ENDPOINT= os.getenv("O4MINI_AZURE_OPENAI_ENDPOINT")
        AZURE_OPENAI_API_KEY = os.getenv("O4MINI_AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_DEPLOYMENT = os.getenv("O4MINI_AZURE_OPENAI_DEPLOYMENT")
        AZURE_OPENAI_VERSION = os.getenv("LATEST_AZURE_OPENAI_VERSION")

        # Initialize the Azure OpenAI client
        if os.getenv("DJANGO_DEVELOPMENT", "False") == "True":
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}",
                http_client=httpx.Client(verify=False)
            )
        else:
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}"
            )

        # Create a chat completion
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            # removing temperature and max_completion changed to max_completion_tokens
            # this is to accommodate to the new change of azureopenai error saying this method
            # is depreciated
            # Removed the max_completion completely when using O1 Models
            # temperature=0.9,
            # max_completion_tokens=2000,
            messages=[{ 
                        "role": "user",
                        "content": f"""
                          {final_prompt}
                        """,
                    }]
        )
        # Extract the generated text from the response
        generated_text = response.choices[0].message.content

        # Return a JSON response
        return Response({"generated_text": generated_text})
class openAICleanVersion_O1MINI(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    
    def _build_prompt(self, data, language, prompt, question, sample):
        """Build the appropriate prompt based on available data"""
        if all(v is None for v in [language, prompt, question, sample]):
            # If only data is provided, send just the data
            return data
        return f"""
        Strictly return only mermaidjs code in plaintext. No explanations or other texts included
        Language: {language}, if none, it is automatic english.
        Also consider architecture-beta of mermaidjs. Sample:
        architecture-beta
            group api(cloud)[API]

            service db(database)[Database] in api
            service disk1(disk)[Storage] in api
            service disk2(disk)[Storage] in api
            service server(server)[Server] in api

            db:L -- R:server
            disk1:T -- B:server
            disk2:T -- B:db
        ---
            {data}
        """


    def post(self, request):
        # Get user input
        data = request.data.get("text")
        language = request.data.get("language")
        prompt = request.data.get("prompt_strategy")
        question = request.data.get("question")
        sample = request.data.get("sample_answer")
        # Load environment variables
        load_dotenv()
        AZURE_OPENAI_ENDPOINT = os.getenv("O1MINI_AZURE_OPENAI_ENDPOINT")
        AZURE_OPENAI_API_KEY = os.getenv("O1MINI_AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_DEPLOYMENT = os.getenv("O1MINI_AZURE_OPENAI_DEPLOYMENT")
        AZURE_OPENAI_VERSION = os.getenv("LATEST_AZURE_OPENAI_VERSION")


        if os.getenv("DJANGO_DEVELOPMENT", "False") == "True":
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}",
                http_client=httpx.Client(verify=False)
            )
        else:
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}"
            )

        # Build the appropriate prompt
        content = self._build_prompt(
            data=data,
            language=language,
            prompt=prompt,
            question=question,
            sample=sample
        )

        # Create a chat completion
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[{
                "role": "user",
                "content": content
            }]
        )
        # Extract the generated text from the response
        generated_text = response.choices[0].message.content
        # Return a JSON response
        return Response({"generated_text": generated_text})

class openAIView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]
    def post(self, request):
        # Get user input
        data = request.data.get("text")
        language = request.data.get("language")
        prompt = request.data.get("prompt_strategy")
        question = request.data.get("question")
        sample = request.data.get("sample_answer")
        # Load environment variables
        load_dotenv()
        AZURE_OPENAI_ENDPOINT= os.getenv("AZURE_OPENAI_ENDPOINT")
        AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
        AZURE_OPENAI_VERSION = os.getenv("LATEST_AZURE_OPENAI_VERSION")

        # Initialize the Azure OpenAI client
        if os.getenv("DJANGO_DEVELOPMENT", "False") == "True":
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}",
                http_client=httpx.Client(verify=False)
            )
        else:
            client = AzureOpenAI(
                api_key=AZURE_OPENAI_API_KEY,
                api_version=AZURE_OPENAI_VERSION,
                base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}"
            )

        # Create a chat completion
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            temperature=0.9,
            max_tokens=2000,
            messages=[
                {
                    "role": "user",
                    "content": f"""
                        For reference, here is the relevant information:
                        Question: {question}
                        Prompt Strategy: {prompt}
                        ---
                        Enhance the provided content while preserving all original information. Improve the text as needed without omitting any details from the initial input. Present only the refined answer in plain text format, 
                        without explanations, questions, formatting, bullet points, headings, or special symbols. Translate all text and diagram contents to the specified language. If unable to follow these instructions, respond 
                        with "We need more information."
                        ---
                        User Input: {data if isinstance(data, str) else data['input_answer']}
                        Translate text to {language}
                        ---
                        Strictly do not include the language, question and prompt strategy in the return answer.
                        If 
                        ---
                        Refine the answer so the following texts are not included in the return text:
                            - {question}
                            - {prompt}
                            - translate {language}
                    """,
                }
            ]
        )
        # Extract the generated text from the response
        generated_text = response.choices[0].message.content
        # Return a JSON response
        return Response({"generated_text": generated_text})
    def get(self, request):
        current_category = request.query_params.get("currentCategory")

        if current_category:
            questions = Question.objects.filter(category=current_category)
        else:
            questions = Question.objects.all()

        serializer = QuestionSerializer(questions, many=True)
        return Response({"question_list": serializer.data}, status=status.HTTP_200_OK)
class ProjectsView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    # List all the projects
    def get(self, request):
        try:
            projects = Project.objects.filter(owner=request.user.id)
            serializer = ProjectSerializer(projects, many=True)
            return Response({"project_list": serializer.data}, status=status.HTTP_200_OK)
        except Project.DoesNotExist:
            return Response({'error': 'Project not found'}, status=status.HTTP_404_NOT_FOUND)
    
    # Create a New Project
    def post(self, request):
        # Create a copy of request.data and add the current user
        data = request.data.copy()  # Make a mutable copy of request.data
        data['owner'] = request.user.id  # Add the current user's ID
        data['owner_name'] = request.user.username
        
        serializer = ProjectSerializer(data=data)  # Pass modified data to serializer

        if serializer.is_valid():
            serializer.save()
            return Response({"message": "Project created successfully", "project": serializer.data}, status=status.HTTP_201_CREATED)
        else:
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    # Edit an existing project
    def put(self, request):
        print("Edit an existing project")

    # Delete a project
    def delete(self, request, pk):
        data = request.data
        pk = pk
        try:
            project = Project.objects.get(pk=pk)
            project.delete()
            return Response({"message": "Project deleted successfully"}, status=status.HTTP_204_NO_CONTENT)
        except Project.DoesNotExist:
            return Response({'error': 'Project not found'}, status=status.HTTP_404_NOT_FOUND)
class VoiceView(APIView):
    def get_gpt_response(request):
        if request.method == 'POST':
            data = json.loads(request.body.decode('utf-8'))
            prompt = data.get('prompt')

            if not prompt:
                return JsonResponse({'error': 'Prompt is required'}, status=400)

            try:
                completion = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                gpt_response = completion.choices[0].message.content
                return JsonResponse({'response': gpt_response})

            except Exception as e:
                return JsonResponse({'error': str(e)}, status=500)
        else:
            return JsonResponse({'error': 'Invalid method'}, status=405)
class AnswersView(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def get(self, request):
        id = request.query_params.get("project_id")
        if id:
            answers = Answer.objects.filter(project_id=id)
        else:
            answers = ""

        serializer = AnswerSerializer(answers, many=True)
        return Response({"answer_list": serializer.data}, status=status.HTTP_200_OK)
    def post(self, request):
        data = request.data        
        # Check if data is a dictionary and not empty
        if isinstance(data, dict) and data:
            answers_created = []  
            for id, answer in data.items():                
                question_id = answer.get('question')
                project_id = answer.get('project_id')
                text = answer.get('input_answer')
                try:
                    question = Question.objects.get(question_id=question_id)

                    # Check if the answer already exists
                    existing_answer = Answer.objects.filter(project_id=project_id, question=question).first()

                    if existing_answer:
                        # Update the existing answer
                        existing_answer.input_answer = text
                        existing_answer.save()
                        answers_created.append(AnswerSerializer(existing_answer).data)
                    else:
                        # Create a new answer
                        answer_data = {
                            "answer_id": f"{project_id}-{question_id}",
                            "project_id": project_id,
                            "question": question.question_id,
                            "input_answer": text,
                            "category": question.category,
                        }
                        serializer = AnswerSerializer(data=answer_data)

                        if serializer.is_valid():
                            serializer.save()
                            answers_created.append(serializer.data)
                        else:
                            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
                except Question.DoesNotExist:
                    return Response({"error": f"Question with id {question_id} does not exist"}, status=status.HTTP_404_NOT_FOUND)

            return Response({"message": "Answers submitted successfully", "answers": answers_created}, status=status.HTTP_201_CREATED)
        
        return Response({"error": "Invalid data format."}, status=status.HTTP_400_BAD_REQUEST)
class GenerateReportView(APIView):
    def get(self, request):
        language = request.query_params.get("language")
        project_id = request.query_params.get("project_id")
        project_answers = Answer.objects.filter(project_id=project_id)
        answers = project_answers.select_related('question').order_by('category', 'question__question_id')
        # doc, answers_text = self.create_document_and_text(answers)
        answers_text = self.create_document_and_text(answers)

        gpt_response = self.generate_summary(answers_text, language)
        message = gpt_response.data.get("generated_text")
        # Create an instance of the DocumentGenerator
        generator = DocumentGenerator()
        # Simulate the response (in a Django view, you would return this response)
        response = generator.create_document_response(message)
        
        return response
    
    def create_document_and_text(self, answers):
        answers_text = ""
        current_category = None
        for answer in answers:
            if answer.category != current_category:
                answers_text += f"\n Section Category: {answer.category}\n"
                current_category = answer.category
            q_text = f"• {answer.question}: {answer.input_answer}\n"
            answers_text += q_text
           
        return answers_text
    
    def generate_summary(self, answers_text, language):
        # Prepare the data for the OpenAIView
        data = {
            "text": answers_text,
            "language": language,
            "prompt_strategy": 
            """
                You are a consultant tasked with creating a comprehensive Risk Advisory Report for your client based on their company data. The report should be professional, detailed, and structured to analyze the provided information without including the specific questions asked of the client.
                Utilize the information from the sections outlined in the "User Input" to craft a cohesive report. The report should:
                - Analyze the Data: Discuss key insights derived from the information provided.
                - Provide Actionable Recommendations: Offer practical steps based on your analysis.
                - Maintain Clarity and Engagement: Use clear and engaging language throughout.
                - Incorporate Visual Aids: Include tables or charts where appropriate to enhance understanding.
                - Ensure Cohesion: Create a narrative that connects insights to practical implications for the client.
                
                ### Formatting Instructions
                Each section header should be formatted as H1 (e.g., ### General Information).
                All other text should be presented in plain text without additional formatting (e.g., no bold or bullet points).
                Strictly follow this order in your final response, only including sections present in the "User Input":
                1. General Information
                2. Authentication and Authorization Concept
                3. Application Design
                4. Cloud Architecture
                5. Architecture Diagram (do not include user input for this section)
                Refer to the "user input" for specific details to include in each section, ensuring that your final report presents a comprehensive overview of the client's risk management strategies and practices. 
                Produce a final report format rather than a simple Q&A structure. Output should be in paragraph forms with proper sections.
                Split sections by double new lines.

                ----
                On each section header measure add a detailed Counter Measure. 
                Countermeasure should include reason and how it should be implemented in deep detail.
                This subsection starts with a header 'Counter Measure' written as H1 (e.g., ### General Information).
                Write in bullet points with 4 lines each.
                ---
                Create a Counter Measure Template with the following specifications:

                1. Include a minimum of 100 words per countermeasure point.
                2. Only include sections where information is available.
                3. Format the response as follows:

                ### Counter Measure

                    1. [Title of first countermeasure]
                        [Detailed explanation of the first countermeasure, including its purpose, implementation, and benefits to security. Ensure this explanation is at least 100 words long.]

                    2. [Title of second countermeasure]
                        [Detailed explanation of the second countermeasure, including its purpose, implementation, and benefits to security. Ensure this explanation is at least 100 words long.]

                    3. [Title of third countermeasure]
                        [Detailed explanation of the third countermeasure, including its purpose, implementation, and benefits to security. Ensure this explanation is at least 100 words long.]"

                Ensure that each countermeasure explanation:
                - Clearly states the purpose of the measure
                - Explains how it enhances security
                - Describes its implementation or application
                - Mentions any potential benefits or improvements to overall system security

                Use a tone that is informative and technical, suitable for a professional cybersecurity context. Number each countermeasure and maintain the given format structure.
            """,
            "question": "Generate a Report",
            "sample_answer": ""
        }

        # Create a mock request with the data
        mock_request = MockRequest(data)

        # Create an instance of openAIView
        open_ai_view = openAIView()

        # Call the post method of openAIView with the mock request
        response = open_ai_view.post(mock_request)
        
        if response.status_code == 200:
            return response
        else:
            return "Failed to generate summary"
class SaveDiagram(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        # Retrieve the diagram data from the request
        diagram_data = request.data.get('diagram', None)
        diagram_name = request.data.get('diagramName')

        if not diagram_data:
            return Response({"error": "No diagram data provided."}, status=status.HTTP_400_BAD_REQUEST)

        # Check if the diagram data is in base64 format for PNG
        if diagram_data.startswith("data:image/png;base64,"):
            # Remove the prefix to get only the base64 string
            diagram_data = diagram_data.replace("data:image/png;base64,", "")
        else:
            return Response({"error": "Invalid image format. Expected PNG."}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Decode the base64 string
            image_data = base64.b64decode(diagram_data)
            file_name = f"{diagram_name}.png"
            file_path = default_storage.path(file_name)  # Get full path of the file

            # Check if the file already exists and delete it
            if default_storage.exists(file_name):
                default_storage.delete(file_name)

            # Save the new image
            path = default_storage.save(file_name, ContentFile(image_data))

            return Response({"message": "Diagram saved successfully.", "path": path}, status=status.HTTP_201_CREATED)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
class TranscribeAudio(APIView):
    authentication_classes = [JWTAuthentication]
    permission_classes = [IsAuthenticated]

    def post(self, request):        
        # Get Base64 audio from request data
        base64_audio = request.data.get("audio")
        audio_format = request.data.get("format")

        if not base64_audio:
            return Response({"error": "Audio data is required."}, status=400)

        # Decode Base64 to binary data
        audio_data = base64.b64decode(base64_audio)

        # Save audio temporarily (optional)
        temp_audio_file = f"temp_audio_{uuid.uuid4().hex}.{audio_format}"
        with open(temp_audio_file, "wb") as f:
            f.write(audio_data)
        temp_audio_path = os.path.abspath(temp_audio_file)
        print()

        try:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            transcript = loop.run_until_complete(transcribe_audio(temp_audio_file))
            return Response({'transcript': transcript})
        
        except Exception as e:
            print("Error", e)
            return Response({'error': str(e)}, status=500)
        finally:
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)
class ProcessDocumentView(APIView):
    permission_classes = [AllowAny]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request):
        file = request.FILES.get('file')
        question = request.data.get('question')
        if not file:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                for chunk in file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            # Read the file content based on its type
            file_extension = os.path.splitext(file.name)[1].lower()
            if file_extension == '.docx':
                doc = docx.Document(temp_file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif file_extension == '.pdf':
                with open(temp_file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
            else:  # Assume it's a text file
                with open(temp_file_path, 'r', encoding='utf-8') as f:
                    text = f.read()

            # Create a mock request with the data
            prompt_text = {"text": f"Answer the question: '{question}' according to the following information: {text} --- If the question is not in the data, return ''"}

            mock_request = MockRequest(prompt_text)
            azure_open_ai = openAICleanVersion()
            # Process the text with Azure OpenAI
            result = azure_open_ai.post(mock_request)
            if result.status_code == 200:
                return result
            else:
                return Response({'error': 'Failed to generate summary'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        finally:
            # Delete the temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

    def put(self, request):
        data = json.loads(request.body)
        ref_answers = data.get('ref_answers')
        project_id = data.get('project_id')
        fill_all = data.get('fill_all', False)

        # Validate input
        if not isinstance(ref_answers, list):
            return Response({"error": "ref_answers must be a list"}, status=status.HTTP_400_BAD_REQUEST)

        project = get_object_or_404(Project, project_id=project_id)
        updated_answer_ids = []

        for ref in ref_answers:
            try:
                question = Question.objects.get(question_id=ref['question_id'])
                try:
                    answer = Answer.objects.create(
                        answer_id =  f"{project_id}-{question.question_id}",
                        question=question,
                        project_id=project,
                        input_answer=ref['ref_answer']
                    )
                except IntegrityError:
                    # Handle the duplicate case, e.g., by updating the existing answer
                    answer = Answer.objects.get(answer_id = f"{project_id}-{question.question_id}")

                isNotEmpty = True if ref.get('ref_answer', '').strip() != '' else False

                if (fill_all or not answer.input_answer) and isNotEmpty:
                    answer.input_answer = ref['ref_answer'].strip()
                    answer.save(update_fields=['input_answer'])
                    updated_answer_ids.append(answer)
                else:
                    if(len(answer.input_answer) < 1):   
                        answer.input_answer = ref['ref_answer'].strip()
                        answer.save(update_fields=['input_answer'])
                        updated_answer_ids.append(answer)
                    else:
                        updated_answer_ids.append(answer)


            except Question.DoesNotExist:
                print(f"Question with id {ref['question_id']} does not exist")
            except KeyError as e:
                print(f"Missing key in ref: {e}")
            serialized_answers = AnswerSerializer(updated_answer_ids, many=True)
            updated_answers_dict = {}
            for answer in serialized_answers.data:
                question_id = answer['question']  # Assuming 'question' field contains the question_id
                updated_answers_dict[question_id] = answer

        return Response({
            "message": "Answers updated successfully",
            "updated_answers": updated_answers_dict
        }, status=status.HTTP_200_OK)

class ExplainImageView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, format=None):
        AZURE_OPENAI_ENDPOINT= os.getenv("AZURE_OPENAI_ENDPOINT")
        AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
        AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
        AZURE_OPENAI_VERSION = os.getenv("LATEST_AZURE_OPENAI_VERSION")

        #PREPROCESSING OF IMAGE
        image = request.FILES.get('image')
        language = request.data.get('language')
        if not image:
            return Response({"detail": "No image uploaded."}, status=status.HTTP_400_BAD_REQUEST)
        if not language:
            return Response({"detail": "No langauge selected."}, status=status.HTTP_400_BAD_REQUEST)
        image.seek(0)
        image_data = base64.b64encode(image.read()).decode('utf-8')

        #CALL GPT
        # Initialize the Azure OpenAI client
        client = AzureOpenAI(
            api_key=AZURE_OPENAI_API_KEY,
            api_version=AZURE_OPENAI_VERSION,
            base_url=f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}",
            http_client=httpx.Client(verify=False)
        )

        # Create a chat completion
        response = client.chat.completions.create(
            model=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                    {"role": "system", "content": "You are a helpful assistant that responds in Markdown. Help me with my math homework!"},
                    {"role": "user", "content": [
                        {"type": "text", "text": "Check if the image given is a diagram. If it is a diagram, do not include 'this is a diagram' Explain this diagram to client in 20 seconds speaking time. If not explain what you see in the image in 5 saying it's not a diagram. Explain in English only"},
                        {"type": "image_url", "image_url": {
                            "url": f"data:image/png;base64,{image_data}"}
                        }   
                    ]}
                ],
            temperature=0.0,
        )
        # Extract the generated text from the response
        generated_text = response.choices[0].message.content
        print("Generated Text Answer:", generated_text)
        # Return a JSON response
        return Response({"explanation": generated_text})
class GenerateDiagramView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        # Get user input from request
        prompt = request.data.get("prompt")
        title = request.data.get("title", "System Architecture")

        gpt_prompt = f"""
            User Input: {prompt}
            Prompt: 
            GPT Prompt for Architecture Diagram Input Generation
            Main Prompt
            You are an expert system architect who converts user requests into detailed architecture descriptions that can be parsed by an automated diagram generator. Your task is to take a user's high-level architecture request and convert it into a comprehensive, detailed description that includes all necessary components and their relationships.
            Instructions
            When a user provides an architecture request, you should:
            Identify the core system type and create a meaningful title
            List all technology components mentioned or implied
            Add typical/standard components that would be expected in such an architecture
            Specify cloud provider (AWS, Azure, GCP) if mentioned or implied
            Include security components that would be standard for the architecture
            Add monitoring and management components as appropriate
            Specify user/client access points
            Describe component relationships and data flow
            Output Format
            Generate a detailed architecture description following this pattern:
            "Create a [SYSTEM_NAME] architecture diagram with [FRONTEND_TECH] as the frontend, [BACKEND_TECH] as the backend, [DATABASE_TECH] database, [SERVICE_DISCOVERY] for service discovery, with typical [CLOUD_PROVIDER] resources. Include [SECURITY_COMPONENTS] for security, [MONITORING_COMPONENTS] for monitoring and management. [ADDITIONAL_SERVICES]. All components inside the [CLOUD_PROVIDER] cloud, with [CLIENT_ACCESS] outside the cloud."
            Component Categories to Consider
            Frontend Technologies
            Angular, React, Vue.js, Web UI, Mobile App, etc.
            Backend Technologies
            Spring Framework, Node.js, .NET Core, Python Flask/Django, Java, etc.
            Databases
            PostgreSQL, MySQL, MongoDB, DynamoDB, CosmosDB, SQL Server, etc.
            Cloud Services (AWS)
            EC2, Lambda, ECS, S3, CloudFront, Route53, ELB, API Gateway, etc.
            Cloud Services (Azure)
            VM, Functions, Storage, SQL Database, Cosmos DB, Application Gateway, Active Directory, etc.
            Cloud Services (GCP)
            Compute Engine, Cloud Functions, Cloud Storage, Cloud SQL, etc.
            Security Components
            IAM, WAF, Firewall, Active Directory, Cognito, Key Vault, etc.
            Monitoring & Management
            CloudWatch, Application Insights, Monitoring, Logging, etc.
            Service Discovery & Integration
            Eureka, Consul, Service Bus, SQS, Event Grid, etc.
            Example Transformations
            Input:
            "Create an Azure architecture diagram illustrating a real-time data processing pipeline. The pipeline should include Azure VM, Azure Storage, Azure SQL, Azure functions and Azure AD. Show the connection of user."
            Output:
            "Create a real-time data processing pipeline architecture diagram with a web dashboard as the frontend, Azure Functions as the serverless backend processing, Azure SQL database for structured data storage, Azure Storage for raw data ingestion, Azure VM for compute-intensive processing tasks, with typical Azure resources. Include Azure AD for authentication and authorization, WAF for firewall protection, Application Gateway for load balancing, Application Insights for monitoring and logging. Additional services include Event Grid for event routing and Service Bus for message queuing. All components inside the Azure cloud, with users and external data sources outside the cloud."
            Key Rules
            Always specify a clear system name/title
            Include both explicitly mentioned and implied standard components
            Add security and monitoring components even if not mentioned
            Specify the cloud provider clearly
            Mention user/client access points
            Include service integration and messaging components for complex systems
            Use specific technology names rather than generic terms
            Ensure logical component relationships
            Additional Guidelines
            If no cloud provider is specified, default to AWS
            If no frontend is mentioned but users are involved, assume a web interface
            Always include basic security (IAM, firewall) and monitoring components
            For data processing pipelines, include appropriate messaging/queuing services
            For microservices, include service discovery and API gateway
            For databases, specify the most appropriate type based on use case
            Response Format
            Provide only the enhanced architecture description without explanations or additional text. The output should be a single, comprehensive sentence that can be directly used as input for the diagram generator.

            Note: 
            - from diagrams.aws.integration import APIGateway gives an error
            """
        
        available_technologies = """
        
        # Frontend technologies
        self.frontend_mapping = {
            'angular': 'diagrams.programming.framework.Angular',
            'angular.js': 'diagrams.programming.framework.Angular',
            'react': 'diagrams.programming.framework.React',
            'vue': 'diagrams.programming.framework.Vue',
            'frontend': 'diagrams.onprem.client.Client',
            'web': 'diagrams.onprem.client.Client',
            'ui': 'diagrams.onprem.client.Client'
        }
        
        # Backend technologies
        self.backend_mapping = {
            'spring': 'diagrams.programming.framework.Spring',
            'javaspring': 'diagrams.programming.framework.Spring',
            'java spring': 'diagrams.programming.framework.Spring',
            'node': 'diagrams.programming.language.Nodejs',
            'nodejs': 'diagrams.programming.language.Nodejs',
            'python': 'diagrams.programming.language.Python',
            'backend': 'diagrams.aws.compute.EC2',
            'api': 'diagrams.aws.compute.EC2'
        }
        
        # Database technologies (multi-cloud)
        self.database_mapping = {
            # Generic/AWS
            'postgresql': 'diagrams.aws.database.RDS',
            'postgres': 'diagrams.aws.database.RDS',
            'mysql': 'diagrams.aws.database.RDS',
            'mongodb': 'diagrams.aws.database.DocumentDB',
            'dynamodb': 'diagrams.aws.database.DynamoDB',
            'redis': 'diagrams.aws.database.ElastiCache',
            'database': 'diagrams.aws.database.RDS',
            # Azure specific
            'azure sql': 'diagrams.azure.database.SQLDatabases',
            'cosmosdb': 'diagrams.azure.database.CosmosDb',
            'cosmos db': 'diagrams.azure.database.CosmosDb',
            'azure cache': 'diagrams.azure.database.CacheForRedis',
            # GCP specific
            'cloud sql': 'diagrams.gcp.database.SQL',
            'firestore': 'diagrams.gcp.database.Firestore',
            'bigtable': 'diagrams.gcp.database.Bigtable',
            'memorystore': 'diagrams.gcp.database.Memorystore'
        }
        
        # AWS services
        self.aws_services = {
            'ec2': 'diagrams.aws.compute.EC2',
            'lambda': 'diagrams.aws.compute.Lambda',
            'ecs': 'diagrams.aws.compute.ECS',
            'eks': 'diagrams.aws.compute.EKS',
            'fargate': 'diagrams.aws.compute.Fargate',
            's3': 'diagrams.aws.storage.S3',
            'cloudfront': 'diagrams.aws.network.CloudFront',
            'route53': 'diagrams.aws.network.Route53',
            'elb': 'diagrams.aws.network.ELB',
            'alb': 'diagrams.aws.network.ALB',
            'vpc': 'diagrams.aws.network.VPC',
            'api gateway': 'diagrams.aws.network.APIGateway',
            'apigateway': 'diagrams.aws.network.APIGateway',
            'gateway': 'diagrams.aws.network.APIGateway',
            'sqs': 'diagrams.aws.integration.SQS',
            'sns': 'diagrams.aws.integration.SNS',
            'eventbridge': 'diagrams.aws.integration.Eventbridge'
        }
        
        # Azure services
        self.azure_services = {
            'vm': 'diagrams.azure.compute.VM',
            'virtual machine': 'diagrams.azure.compute.VM',
            'azure vm': 'diagrams.azure.compute.VM',
            'app service': 'diagrams.azure.compute.AppServices',
            'function app': 'diagrams.azure.compute.FunctionApps',
            'azure functions': 'diagrams.azure.compute.FunctionApps',
            'aks': 'diagrams.azure.compute.KubernetesServices',
            'container instances': 'diagrams.azure.compute.ContainerInstances',
            'storage account': 'diagrams.azure.storage.StorageAccounts',
            'azure storage': 'diagrams.azure.storage.StorageAccounts',
            'blob storage': 'diagrams.azure.storage.BlobStorage',
            'cdn': 'diagrams.azure.network.CDNProfiles',
            'application gateway': 'diagrams.azure.network.ApplicationGateway',
            'load balancer': 'diagrams.azure.network.LoadBalancers',
            'vnet': 'diagrams.azure.network.VirtualNetworks',
            'virtual network': 'diagrams.azure.network.VirtualNetworks',
            'api management': 'diagrams.azure.integration.APIManagement',
            'service bus': 'diagrams.azure.integration.ServiceBus',
            'event grid': 'diagrams.azure.integration.IntegrationAccounts',
            'event hub': 'diagrams.azure.analytics.EventHubs',
            'logic apps': 'diagrams.azure.integration.LogicApps'
        }
        
        # GCP services
        self.gcp_services = {
            'compute engine': 'diagrams.gcp.compute.ComputeEngine',
            'gce': 'diagrams.gcp.compute.ComputeEngine',
            'cloud functions': 'diagrams.gcp.compute.Functions',
            'cloud run': 'diagrams.gcp.compute.Run',
            'gke': 'diagrams.gcp.compute.GKE',
            'kubernetes engine': 'diagrams.gcp.compute.GKE',
            'app engine': 'diagrams.gcp.compute.AppEngine',
            'cloud storage': 'diagrams.gcp.storage.Storage',
            'gcs': 'diagrams.gcp.storage.Storage',
            'cloud cdn': 'diagrams.gcp.network.CDN',
            'load balancing': 'diagrams.gcp.network.LoadBalancing',
            'vpc': 'diagrams.gcp.network.VPC',
            'api gateway': 'diagrams.oci.devops.APIGateway',
            'pub/sub': 'diagrams.gcp.analytics.PubSub',
            'pubsub': 'diagrams.gcp.analytics.PubSub',
            'cloud scheduler': 'diagrams.gcp.devtools.Scheduler',
            'cloud tasks': 'diagrams.gcp.compute.Functions'  # Using Functions as placeholder
        }
        
        # Security components (multi-cloud)
        self.security_mapping = {
            # AWS
            'iam': 'diagrams.aws.security.IAM',
            'waf': 'diagrams.aws.security.WAF',
            'shield': 'diagrams.aws.security.Shield',
            'cognito': 'diagrams.aws.security.Cognito',
            # Azure
            'azure ad': 'diagrams.azure.identity.ActiveDirectory',
            'active directory': 'diagrams.azure.identity.ActiveDirectory',
            'key vault': 'diagrams.azure.security.KeyVaults',
            'azure firewall': 'diagrams.azure.network.Firewall',
            # GCP
            'cloud iam': 'diagrams.gcp.security.IAM',
            'cloud armor': 'diagrams.gcp.security.Armor',
            'cloud kms': 'diagrams.gcp.security.KMS',
            # Generic
            'firewall': 'diagrams.aws.security.WAF',
            'identity': 'diagrams.aws.security.IAM'
        }
        
        # Monitoring and management (multi-cloud)
        self.monitoring_mapping = {
            # AWS
            'cloudwatch': 'diagrams.aws.management.Cloudwatch',
            'cloudtrail': 'diagrams.aws.management.Cloudtrail',
            # Azure
            'azure monitor': 'diagrams.azure.devops.Monitor',
            'application insights': 'diagrams.azure.devops.ApplicationInsights',
            'log analytics': 'diagrams.azure.analytics.LogAnalytics',
            # GCP
            'cloud monitoring': 'diagrams.gcp.operations.Monitoring',
            'cloud logging': 'diagrams.gcp.operations.Logging',
            'cloud trace': 'diagrams.gcp.operations.Trace',
            # Generic
            'monitoring': 'diagrams.aws.management.Cloudwatch',
            'logging': 'diagrams.aws.management.Cloudtrail'
        }
        
        # Container and orchestration (multi-cloud)
        self.container_mapping = {
            # Generic
            'docker': 'diagrams.onprem.container.Docker',
            'kubernetes': 'diagrams.k8s.compute.Pod',
            'k8s': 'diagrams.k8s.compute.Pod',
            # AWS
            'ecs': 'diagrams.aws.compute.ECS',
            'eks': 'diagrams.aws.compute.EKS',
            'fargate': 'diagrams.aws.compute.Fargate',
            # Azure
            'aks': 'diagrams.azure.compute.KubernetesServices',
            'container instances': 'diagrams.azure.compute.ContainerInstances',
            # GCP
            'gke': 'diagrams.gcp.compute.GKE',
            'cloud run': 'diagrams.gcp.compute.Run'
        }
        
        # Service discovery and messaging
        self.service_discovery = {
            # Generic
            'eureka': 'diagrams.aws.integration.SQS',
            'consul': 'diagrams.aws.integration.SQS',
            'service discovery': 'diagrams.aws.integration.SQS',
            'service exploration': 'diagrams.aws.integration.SQS',
            # AWS
            'sqs': 'diagrams.aws.integration.SQS',
            'sns': 'diagrams.aws.integration.SNS',
            # Azure
            'service bus': 'diagrams.azure.integration.ServiceBus',
            'event grid': 'diagrams.azure.integration.IntegrationAccounts',
            # GCP
            'pub/sub': 'diagrams.gcp.analytics.PubSub',
            'pubsub': 'diagrams.gcp.analytics.PubSub'
        }
        
        # User/Client
        self.client_mapping = {
            'user': 'diagrams.onprem.client.Users',
            'users': 'diagrams.onprem.client.Users',
            'client': 'diagrams.onprem.client.Client'
        }
        
        # DevOps and CI/CD (multi-cloud)
        self.devops_mapping = {
            # Generic
            'jenkins': 'diagrams.onprem.ci.Jenkins',
            'gitlab': 'diagrams.onprem.vcs.Gitlab',
            'github': 'diagrams.onprem.vcs.Github',
            # AWS
            'codebuild': 'diagrams.aws.devtools.Codebuild',
            'codepipeline': 'diagrams.aws.devtools.Codepipeline',
            'codecommit': 'diagrams.aws.devtools.Codecommit',
            # Azure
            'azure devops': 'diagrams.azure.devops.Devops',
            'azure pipelines': 'diagrams.azure.devops.Pipelines',
            'azure repos': 'diagrams.azure.devops.Repos',
            # GCP
            'cloud build': 'diagrams.gcp.devtools.Build',
            'cloud source repositories': 'diagrams.gcp.devtools.SourceRepositories'
        }
        """
        print("PROMPT ENTERED", prompt)
        if not prompt:
            return Response(
                {"error": "Prompt is required"},
                status=status.HTTP_400_BAD_REQUEST
            )

        gpt = ClassicGPT()
        max_retries = 3
        current_prompt = gpt_prompt  # Start with initial prompt

        for attempt in range(max_retries):
            print(f"Attempt {attempt+1} with prompt: {current_prompt}")
            result = gpt.getGPTResponse(current_prompt)
            print("FINAL RESULT", result)
            
            # Generate diagram from GPT response
            diagram_result = generate_architecture_diagram(result, title)
            print("CODE:", diagram_result['code'] )

            
            if diagram_result['success']:
                # Process successful diagram
                image_path = diagram_result['image_path']
                print("Image Path:", image_path)
                if not os.path.exists(image_path):
                    return Response(
                        {"error": "Generated image not found"},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
                try:
                    with open(image_path, 'rb') as image_file:
                        image_data = image_file.read()
                    os.unlink(image_path)
                    return HttpResponse(image_data, content_type='image/png', status=200)
                except Exception as e:
                    return Response(
                        {"error": f"File handling error: {str(e)}"},
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
            else:
                # On failure: Extract error and regenerate prompt
                error_message = diagram_result.get('error', 'Diagram generation failed')
                print(f"Attempt {attempt+1} failed: {error_message}")
                
                # Append error to prompt for regeneration
                current_prompt = (
                    f"this is the code:"
                    f"{diagram_result['code']}\n\n"
                    f"Error during execution: {error_message}\n"
                    f"Please analyze the error and regenerate valid diagram code. Strictly only use what is in {available_technologies} "
                )

        # All retries failed
        return Response(
            {"error": "Failed to generate diagram after 3 attempts"},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )